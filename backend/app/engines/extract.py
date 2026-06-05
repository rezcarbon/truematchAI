"""Document text extraction for resume intake.

Converts an uploaded resume (PDF / DOCX / plain text) into raw text that the
intake engine can parse. Extraction runs at upload time so the parsed text is
persisted with the resume record and the assessment pipeline never re-reads the
binary object.

Dependencies: `pypdf` (PDF) and `python-docx` (DOCX). Both are pure-Python and
import lazily so a missing optional dependency only affects that one format.
"""
from __future__ import annotations

import io
import logging

logger = logging.getLogger("truematch.extract")

# Guard against pathological inputs (e.g. a 500-page scanned PDF) producing an
# unbounded prompt. Resumes are short; anything beyond this is truncated.
MAX_CHARS = 60_000


class ExtractionError(RuntimeError):
    """Raised when text cannot be extracted from an uploaded document."""


def extract_text(data: bytes, file_type: str) -> str:
    """Extract UTF-8 text from a resume document.

    file_type: one of "pdf", "docx", "txt".
    Returns normalized text (whitespace-collapsed, length-bounded).
    """
    ft = (file_type or "").lower()
    if ft == "pdf":
        text = _extract_pdf(data)
    elif ft == "docx":
        text = _extract_docx(data)
    elif ft in {"txt", "text", "plain"}:
        text = data.decode("utf-8", errors="replace")
    else:
        raise ExtractionError(f"Unsupported document type for extraction: {file_type!r}")

    text = _normalize(text)
    if not text:
        raise ExtractionError(
            "No extractable text found (the document may be image-only and require OCR)."
        )
    return text[:MAX_CHARS]


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # lazy import
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ExtractionError("PDF extraction requires the 'pypdf' package.") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:  # noqa: BLE001 - pypdf raises a variety of types
        raise ExtractionError(f"Failed to read PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx, lazy import
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ExtractionError("DOCX extraction requires the 'python-docx' package.") from exc
    try:
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        # Include table cell text — resumes frequently use tables for layout.
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        # Fallback: try to extract text directly from the ZIP structure
        # This handles corrupted DOCX files that python-docx can't read
        try:
            text = _extract_docx_fallback(data)
            if text:
                return text
        except Exception:  # noqa: BLE001
            pass
        raise ExtractionError(f"Failed to read DOCX: {exc}") from exc


def _extract_docx_fallback(data: bytes) -> str:
    """Fallback DOCX text extraction from raw ZIP structure.

    Handles corrupted or non-standard DOCX files by extracting text
    directly from the document.xml file without using python-docx.
    """
    import zipfile
    import xml.etree.ElementTree as ET

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as docx_zip:
            # Try to read document.xml
            try:
                xml_content = docx_zip.read('word/document.xml')
            except KeyError:
                return ""

            # Parse XML and extract all text nodes
            root = ET.fromstring(xml_content)
            namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Find all text elements
            text_elements = root.findall('.//w:t', namespace)
            text = "".join([elem.text or "" for elem in text_elements])

            return text if text.strip() else ""
    except Exception:  # noqa: BLE001
        return ""


def _normalize(text: str) -> str:
    # Collapse runs of blank lines and trailing spaces while preserving structure.
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").split("\n")]
    cleaned: list[str] = []
    blank = False
    for line in lines:
        if line.strip():
            cleaned.append(line)
            blank = False
        elif not blank:
            cleaned.append("")
            blank = True
    return "\n".join(cleaned).strip()
